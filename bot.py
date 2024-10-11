import os
import logging
import asyncio
import zipfile
from dotenv import load_dotenv
from pydub import AudioSegment
from docx import Document
import numpy as np
import librosa
import soundfile as sf
import noisereduce as nr
from scipy.signal import butter, filtfilt

import openai
from openai import OpenAIError  # Import corretto secondo la tua indicazione
from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ConversationHandler,
    ContextTypes,
)

# Configura il logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Carica le variabili d'ambiente dal file .env
load_dotenv()

# Stati per le conversazioni
# Per la trascrizione
TRANS_WAITING_FOR_AUDIO, TRANS_WAITING_FOR_FILENAME = range(2)
# Per la pulizia
CLEAN_WAITING_FOR_AUDIO, CLEAN_WAITING_FOR_FILENAME = range(2)

# Variabili OpenAI
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    logger.error("La chiave API di OpenAI non è stata trovata nel file .env.")
    exit(1)

# Configura OpenAI
openai.api_key = OPENAI_API_KEY

# Funzione per il comando /start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        with open('welcome.txt', 'r', encoding='utf-8') as f:
            welcome_message = f.read()
    except Exception as e:
        logger.error(f"Errore durante la lettura del messaggio di benvenuto: {e}")
        welcome_message = "Benvenuto nel bot multifunzionale! Si è verificato un errore nel caricamento del messaggio di benvenuto."

    await update.message.reply_text(welcome_message, parse_mode='Markdown')

######################
# Trascrizione Audio #
######################

# Funzione per il comando /transcribe
async def transcribe_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Per favore, inviami il file audio che desideri trascrivere.")
    return TRANS_WAITING_FOR_AUDIO

# Funzione per ricevere il file audio nella trascrizione
async def transcribe_handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        # Supporta diversi tipi di file audio (audio o voice)
        if update.message.audio:
            audio_file = await update.message.audio.get_file()
        elif update.message.voice:
            audio_file = await update.message.voice.get_file()
        else:
            await update.message.reply_text("Per favore, inviami un file audio valido.")
            return ConversationHandler.END

        file_extension = os.path.splitext(audio_file.file_path)[1]
        if not file_extension:
            file_extension = ".ogg"  # Default a .ogg se non c'è estensione

        file_path = f'tmp/transcribe_received_audio{file_extension}'
        os.makedirs('tmp', exist_ok=True)
        await audio_file.download_to_drive(file_path)

        await update.message.reply_text("File ricevuto! Ora, per favore, inserisci il nome che vuoi dare al file di trascrizione.")
        context.user_data['transcribe_audio_file_path'] = file_path
        return TRANS_WAITING_FOR_FILENAME

    except Exception as e:
        logger.error(f"Errore durante la ricezione dell'audio per trascrizione: {e}")
        await update.message.reply_text("C'è stato un errore durante la ricezione del tuo audio. Per favore, riprova.")
        return ConversationHandler.END

# Funzione per ricevere il nome del file e avviare la trascrizione
async def transcribe_receive_filename(update: Update, context: ContextTypes.DEFAULT_TYPE):
    filename = update.message.text.strip()
    if not filename:
        await update.message.reply_text("Il nome del file non può essere vuoto. Per favore, inserisci un nome valido.")
        return TRANS_WAITING_FOR_FILENAME

    context.user_data['transcribe_filename'] = filename
    await update.message.reply_text(f"Nome del file impostato: `{filename}.docx`. Procedo con la trascrizione dell'audio...", parse_mode='Markdown')

    # Inizia la trascrizione dell'audio
    transcription = await transcribe_audio(update, context)

    if transcription:
        if isinstance(transcription, list):
            # Se transcription è una lista, significa che l'audio è stato suddiviso e trascritto in più parti
            zip_path = os.path.join("tmp", f"{context.user_data['transcribe_filename']}.zip")
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for idx, text in enumerate(transcription, 1):
                    doc_name = f"{context.user_data['transcribe_filename']}_part{idx}.docx"
                    doc_path = os.path.join("tmp", doc_name)
                    create_docx(text, doc_path)
                    zipf.write(doc_path, arcname=doc_name)
                    os.remove(doc_path)  # Rimuovi il .docx dopo averlo aggiunto allo zip

            # Invia il file .zip all'utente
            with open(zip_path, 'rb') as zip_file:
                await update.message.reply_document(document=zip_file, filename=f"{context.user_data['transcribe_filename']}.zip")

            os.remove(zip_path)  # Rimuovi lo zip dopo l'invio
        else:
            # Se transcription è una stringa, invia il .docx come singolo file
            if transcription.strip() == "":
                await update.message.reply_text("La trascrizione è vuota. Assicurati che l'audio contenga parlato chiaro.")
            else:
                # Crea il file .docx con la trascrizione
                doc_path = os.path.join("tmp", f"{context.user_data['transcribe_filename']}.docx")
                create_docx(transcription, doc_path)

                # Invia il file .docx all'utente
                with open(doc_path, 'rb') as doc_file:
                    await update.message.reply_document(document=doc_file, filename=context.user_data['transcribe_filename'] + ".docx")

                # Pulizia dei file temporanei
                os.remove(doc_path)
    else:
        await update.message.reply_text("C'è stato un errore durante la trascrizione del tuo audio. Per favore, riprova.")

    # Pulizia del file audio originale
    file_path = context.user_data.get('transcribe_audio_file_path')
    if file_path and os.path.exists(file_path):
        os.remove(file_path)

    return ConversationHandler.END

# Funzione per creare un file .docx con la trascrizione
def create_docx(transcription_text, doc_path):
    try:
        document = Document()
        document.add_heading('Trascrizione Audio', 0)
        document.add_paragraph(transcription_text)
        document.save(doc_path)
        logger.info(f"Documento .docx creato: {doc_path}")
    except Exception as e:
        logger.error(f"Errore durante la creazione del file .docx: {e}")
        raise

# Funzione per suddividere l'audio in chunk di 15 minuti
def split_audio(file_path, chunk_length_ms=15*60*1000):
    try:
        audio = AudioSegment.from_file(file_path)
        chunks = []
        for i in range(0, len(audio), chunk_length_ms):
            chunk = audio[i:i+chunk_length_ms]
            chunk_filename = f"{file_path.rsplit('.', 1)[0]}_chunk{i//chunk_length_ms +1}.wav"
            chunk.export(chunk_filename, format="wav")
            chunks.append(chunk_filename)
            logger.info(f"Audio suddiviso in: {chunk_filename}")
        return chunks
    except Exception as e:
        logger.error(f"Errore durante la suddivisione dell'audio: {e}")
        return []

# Funzione asincrona per gestire la trascrizione usando Whisper
async def transcribe_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file_path = context.user_data['transcribe_audio_file_path']
    try:
        # Converti il file audio in WAV se necessario
        audio = AudioSegment.from_file(file_path)
        wav_path = file_path.rsplit('.', 1)[0] + ".wav"
        audio.export(wav_path, format="wav")
        logger.info(f"Audio convertito in WAV: {wav_path}")

        # Funzione sincrona per la trascrizione
        def transcribe():
            with open(wav_path, 'rb') as audio_file:
                try:
                    # Richiedi la trascrizione a OpenAI Whisper
                    transcript = openai.Audio.transcribe(
                        model="whisper-1",
                        file=audio_file,
                        language="it"
                    )
                    return transcript.get('text', "").strip()
                except OpenAIError as e:
                    logger.error(f"Errore durante la trascrizione: {e}")
                    if "file too large" in str(e).lower():
                        return "FILE_TOO_LARGE"
                    return None
                except Exception as e:
                    logger.error(f"Errore durante la trascrizione: {e}")
                    return None

        # Esegui la trascrizione in un thread separato
        loop = asyncio.get_event_loop()
        transcription_text = await loop.run_in_executor(None, transcribe)

        if transcription_text == "FILE_TOO_LARGE":
            # Gestisci la suddivisione dell'audio
            await update.message.reply_text(
                "Il file audio inviato è troppo grande per essere trascritto in un'unica parte. "
                "Suddividerò l'audio in parti più piccole e trascriverò ciascuna parte separatamente."
            )
            # Suddividi l'audio in chunk
            chunks = split_audio(wav_path)
            if not chunks:
                await update.message.reply_text("C'è stato un errore durante la suddivisione dell'audio. Per favore, riprova.")
                return None

            transcriptions = []
            for chunk_path in chunks:
                transcription = await loop.run_in_executor(None, transcribe_chunk, chunk_path)
                if transcription:
                    transcriptions.append(transcription)
                else:
                    transcriptions.append("Errore durante la trascrizione di questa parte.")

            # Rimuovi i chunk temporanei
            for chunk_path in chunks:
                if os.path.exists(chunk_path):
                    os.remove(chunk_path)
                    logger.info(f"File chunk rimosso: {chunk_path}")

            # Rimuovi il file WAV temporaneo
            if os.path.exists(wav_path):
                os.remove(wav_path)
                logger.info(f"File WAV temporaneo rimosso: {wav_path}")

            return transcriptions
        elif transcription_text:
            logger.info(f"Trascrizione ottenuta: {transcription_text[:50]}...")
            return transcription_text
        else:
            return None

    except Exception as e:
        logger.error(f"Errore durante la trascrizione dell'audio: {e}")
        return None

def transcribe_chunk(chunk_path):
    try:
        with open(chunk_path, 'rb') as chunk_file:
            chunk_transcript = openai.Audio.transcribe(
                model="whisper-1",
                file=chunk_file,
                language="it"
            )
            return chunk_transcript.get('text', "").strip()
    except OpenAIError as e:
        logger.error(f"Errore durante la trascrizione della parte {chunk_path}: {e}")
        if "file too large" in str(e).lower():
            return "FILE_TOO_LARGE"
        return None
    except Exception as e:
        logger.error(f"Errore durante la trascrizione della parte {chunk_path}: {e}")
        return None

##################
# Pulizia Audio  #
##################

# Funzione per il comando /clean
async def clean_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Per favore, inviami il file audio che desideri pulire.")
    return CLEAN_WAITING_FOR_AUDIO

# Funzione per ricevere il file audio nella pulizia
async def clean_handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        if update.message.audio:
            audio_file = await update.message.audio.get_file()
        elif update.message.voice:
            audio_file = await update.message.voice.get_file()
        else:
            await update.message.reply_text("Per favore, inviami un file audio valido.")
            return ConversationHandler.END

        file_extension = os.path.splitext(audio_file.file_path)[1]
        if not file_extension:
            file_extension = ".ogg"  # Default a .ogg se non c'è estensione

        file_path = f'tmp/clean_received_audio{file_extension}'
        os.makedirs('tmp', exist_ok=True)
        await audio_file.download_to_drive(file_path)

        await update.message.reply_text("File ricevuto! Ora, per favore, inserisci il nome che vuoi dare al file pulito.")
        context.user_data['clean_audio_file_path'] = file_path
        return CLEAN_WAITING_FOR_FILENAME

    except Exception as e:
        logger.error(f"Errore durante la ricezione dell'audio per pulizia: {e}")
        await update.message.reply_text("C'è stato un errore durante la ricezione del tuo audio. Per favore, riprova.")
        return ConversationHandler.END

# Funzione per ricevere il nome del file e avviare la pulizia
async def clean_receive_filename(update: Update, context: ContextTypes.DEFAULT_TYPE):
    filename = update.message.text.strip()
    if not filename:
        await update.message.reply_text("Il nome del file non può essere vuoto. Per favore, inserisci un nome valido.")
        return CLEAN_WAITING_FOR_FILENAME

    context.user_data['clean_filename'] = filename + ".mp3"
    await update.message.reply_text(f"Nome del file impostato: `{filename}.mp3`. Procedo con la pulizia dell'audio...", parse_mode='Markdown')

    # Inizia la pulizia dell'audio
    await clean_audio(update, context)

    return ConversationHandler.END

# Funzione per applicare filtri Butterworth
def butter_filter(data, lowcut, highcut, fs, btype='low'):
    nyq = 0.5 * fs
    low = lowcut / nyq
    high = highcut / nyq
    if btype == 'band':
        b, a = butter(4, [low, high], btype='band')
    else:
        b, a = butter(4, low if btype == 'low' else high, btype=btype)
    y = filtfilt(b, a, data)
    return y

# Funzione per pulire l'audio
async def clean_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file_path = context.user_data['clean_audio_file_path']
    filename = context.user_data['clean_filename']
    
    try:
        # Carica il file audio
        audio = AudioSegment.from_file(file_path)

        if len(audio) == 0:
            await update.message.reply_text("L'audio è vuoto. Non posso pulirlo.")
            return

        # Normalizzazione del volume
        normalized_audio = audio.normalize()

        # Converti in un array NumPy
        y = np.array(normalized_audio.get_array_of_samples(), dtype=np.float32)
        y /= np.max(np.abs(y))  # Normalizza

        # Riduzione del rumore
        noise_factor = 0.2  # Prova diversi valori tra 0.1 e 0.5
        y_denoised = nr.reduce_noise(y=y, sr=normalized_audio.frame_rate, prop_decrease=noise_factor)

        # Equalizzazione del segnale
        low_cutoff_freq = 300.0  # Frequenza di taglio inferiore in Hz
        high_cutoff_freq = 3400.0  # Frequenza di taglio superiore in Hz
        
        # Filtro high-pass
        y_filtered_hp = butter_filter(y_denoised, low_cutoff_freq, high_cutoff_freq, normalized_audio.frame_rate, btype='high')
        
        # Filtro low-pass
        y_filtered_lp = butter_filter(y_filtered_hp, low_cutoff_freq, high_cutoff_freq, normalized_audio.frame_rate, btype='low')

        # Salva l'audio pulito in un file MP3
        mp3_path = os.path.join("tmp", filename)
        sf.write(mp3_path, y_filtered_lp, normalized_audio.frame_rate)
        logger.info(f"Audio pulito salvato: {mp3_path}")

        # Invia l'audio pulito all'utente
        with open(mp3_path, 'rb') as audio_file_to_send:
            await update.message.reply_audio(audio=audio_file_to_send)

        # Pulizia dei file temporanei
        os.remove(mp3_path)
        os.remove(file_path)
        logger.info(f"File temporanei rimossi: {mp3_path}, {file_path}")

    except Exception as e:
        logger.error(f"Errore durante la pulizia dell'audio: {e}")
        await update.message.reply_text("C'è stato un errore durante la pulizia del tuo audio. Per favore, riprova.")

##########################
# Gestione delle Conversazioni #
##########################

# Funzione per annullare la conversazione
async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Operazione annullata.')
    return ConversationHandler.END

# Funzione di gestione degli errori
async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.error(msg="Eccezione non gestita:", exc_info=context.error)
    if isinstance(update, Update) and update.message:
        await update.message.reply_text("Si è verificato un errore. Per favore, riprova più tardi.")

def main():
    TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
    if not TOKEN:
        logger.error("Il token del bot non è stato trovato nel file .env.")
        return

    application = Application.builder().token(TOKEN).build()

    # Handler per la trascrizione
    transcribe_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("transcribe", transcribe_command)],
        states={
            TRANS_WAITING_FOR_AUDIO: [MessageHandler(filters.AUDIO | filters.VOICE, transcribe_handle_audio)],
            TRANS_WAITING_FOR_FILENAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, transcribe_receive_filename)],
        },
        fallbacks=[CommandHandler("cancel", cancel)]
    )

    # Handler per la pulizia
    clean_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("clean", clean_command)],
        states={
            CLEAN_WAITING_FOR_AUDIO: [MessageHandler(filters.AUDIO | filters.VOICE, clean_handle_audio)],
            CLEAN_WAITING_FOR_FILENAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, clean_receive_filename)],
        },
        fallbacks=[CommandHandler("cancel", cancel)]
    )

    # Aggiungi l'handler per il comando /start
    application.add_handler(CommandHandler("start", start))

    # Aggiungi gli handler per le conversazioni
    application.add_handler(transcribe_conv_handler)
    application.add_handler(clean_conv_handler)

    # Aggiungi l'handler per gli errori
    application.add_error_handler(error_handler)

    logger.info("Avvio del bot...")
    application.run_polling()

if __name__ == '__main__':
    main()

